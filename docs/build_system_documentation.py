from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "TAI-LIEU-KY-THUAT-HE-THONG-SAVING-100.md"
OUTPUT = ROOT / "TAI-LIEU-KY-THUAT-HE-THONG-SAVING-100.docx"
ASSETS = ROOT / "assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CAD3DF"
WHITE = "FFFFFF"
BLACK = "111111"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def font_path(bold: bool = False) -> str:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path(bold), size=size)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    fill: str = "#FFFFFF",
    outline: str = "#CAD3DF",
    accent: str = "#2E74B5",
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    draw.rounded_rectangle((x1, y1, x2, y1 + 58), radius=22, fill=accent)
    draw.rectangle((x1, y1 + 34, x2, y1 + 58), fill=accent)
    title_font = pil_font(26, True)
    body_font = pil_font(21)
    draw.text((x1 + 20, y1 + 14), title, font=title_font, fill="#FFFFFF")
    y = y1 + 75
    for line in lines:
        draw.text((x1 + 20, y), line, font=body_font, fill="#283548")
        y += 31


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
    color: str = "#6B7A90",
) -> None:
    draw.line((start, end), fill=color, width=5)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 16
    p1 = (x2 - ux * size + px * size * 0.6, y2 - uy * size + py * size * 0.6)
    p2 = (x2 - ux * size - px * size * 0.6, y2 - uy * size - py * size * 0.6)
    draw.polygon([end, p1, p2], fill=color)
    if label:
        font = pil_font(18, True)
        tx = int((x1 + x2) / 2)
        ty = int((y1 + y2) / 2) - 28
        bbox = draw.textbbox((0, 0), label, font=font)
        width = bbox[2] - bbox[0]
        draw.rounded_rectangle((tx - width // 2 - 8, ty - 3, tx + width // 2 + 8, ty + 24), radius=8, fill="#FFFFFF")
        draw.text((tx - width // 2, ty), label, font=font, fill=color)


def poly_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    label: str = "",
    label_at: tuple[int, int] | None = None,
    color: str = "#6B7A90",
) -> None:
    draw.line(points, fill=color, width=5, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 16
    p1 = (x2 - ux * size + px * size * 0.6, y2 - uy * size + py * size * 0.6)
    p2 = (x2 - ux * size - px * size * 0.6, y2 - uy * size - py * size * 0.6)
    draw.polygon([(x2, y2), p1, p2], fill=color)
    if label:
        font = pil_font(18, True)
        tx, ty = label_at if label_at else points[len(points) // 2]
        bbox = draw.textbbox((0, 0), label, font=font)
        width = bbox[2] - bbox[0]
        draw.rounded_rectangle((tx - width // 2 - 8, ty - 3, tx + width // 2 + 8, ty + 24), radius=8, fill="#FFFFFF")
        draw.text((tx - width // 2, ty), label, font=font, fill=color)


def create_architecture_image(path: Path) -> None:
    image = Image.new("RGB", (1800, 1050), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "KIẾN TRÚC TỔNG THỂ SAVING 100", font=pil_font(42, True), fill="#0B2545")
    draw.text((70, 102), "Next.js + NestJS + MongoDB Atlas + payOS + Web Push", font=pil_font(25), fill="#667085")

    rounded_box(draw, (80, 235, 470, 550), "Frontend / PWA", [
        "Next.js App Router",
        "React Query + Zustand",
        "AuthProvider / AuthGate",
        "Service Worker / QR UI",
    ], accent="#0B2545")
    rounded_box(draw, (700, 205, 1110, 590), "NestJS API", [
        "/api/v1 REST",
        "Auth / Challenge / Plan",
        "Payment / Push / Health",
        "Validation / CORS / Helmet",
        "Cron scheduler",
    ], accent="#2E74B5")
    rounded_box(draw, (1340, 240, 1720, 540), "MongoDB Atlas", [
        "12 collections",
        "Mongoose schemas",
        "Indexes / validators",
        "Database: saving_100_app",
    ], accent="#1F6F5B")

    rounded_box(draw, (630, 735, 970, 970), "payOS", [
        "Payment link / QR",
        "Reconcile / cancel",
        "Signed webhook",
    ], accent="#7A5A00")
    rounded_box(draw, (1130, 735, 1700, 970), "Web Push Network", [
        "VAPID encryption",
        "Google / Apple / Mozilla endpoint",
        "Push về trình duyệt người dùng",
    ], accent="#6C4BB6")

    arrow(draw, (470, 385), (700, 385), "HTTPS + JWT")
    arrow(draw, (1110, 380), (1340, 380), "Mongoose/TLS")
    arrow(draw, (790, 590), (750, 735), "SDK")
    arrow(draw, (900, 735), (950, 590), "Webhook")
    arrow(draw, (1035, 590), (1250, 735), "VAPID")
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=95)


def create_erd_image(path: Path) -> None:
    image = Image.new("RGB", (2100, 1500), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "ERD LOGIC HIỆN TẠI", font=pil_font(44, True), fill="#0B2545")
    draw.text((60, 92), "Đường nối thể hiện quan hệ ObjectId/logic; MongoDB không cưỡng chế foreign key.", font=pil_font(23), fill="#667085")

    boxes = {
        "users": (80, 190, 440, 410),
        "push": (80, 520, 440, 750),
        "challenges": (590, 190, 980, 420),
        "checkins": (540, 540, 900, 790),
        "challenge_payments": (940, 540, 1325, 815),
        "plans": (1450, 180, 2010, 455),
        "slots": (1390, 560, 1750, 825),
        "payments": (1790, 560, 2070, 865),
        "records": (1390, 960, 1750, 1220),
        "events": (950, 980, 1270, 1240),
        "webhooks": (500, 1030, 880, 1290),
        "counters": (80, 1040, 420, 1270),
    }

    rounded_box(draw, boxes["users"], "users", ["_id", "email / passwordHash", "refreshTokenHash", "timezone / push settings"], accent="#0B2545")
    rounded_box(draw, boxes["push"], "push_subscriptions", ["userId", "endpoint (unique)", "keys", "lastUsedAt"], accent="#6C4BB6")
    rounded_box(draw, boxes["challenges"], "saving_challenges", ["userId", "1..100 / unitAmount", "savedAmount", "status"], accent="#2E74B5")
    rounded_box(draw, boxes["checkins"], "saving_checkins", ["challengeId / userId", "number / amount", "localDate", "COMPLETED / REVERSED"], accent="#2E74B5")
    rounded_box(draw, boxes["challenge_payments"], "challenge_payments", ["challengeId / number", "orderCode", "QR / checkoutUrl", "payment status"], accent="#7A5A00")
    rounded_box(draw, boxes["plans"], "saving_plans", ["userId / durationDays", "generationMode", "aggregate progress", "status / confirmationMode"], accent="#1F6F5B")
    rounded_box(draw, boxes["slots"], "saving_slots", ["planId / slotIndex", "amount", "reservation", "slot status"], accent="#1F6F5B")
    rounded_box(draw, boxes["payments"], "saving_payments", ["planId / slotId", "orderCode", "QR / checkoutUrl", "payment status"], accent="#7A5A00")
    rounded_box(draw, boxes["records"], "saving_day_records", ["planId / slotId", "paymentId?", "dayIndex / amount", "PAYOS / MANUAL"], accent="#1F6F5B")
    rounded_box(draw, boxes["events"], "saving_events", ["planId / optional refs", "event type", "before / after", "dormant"], accent="#8B5E83")
    rounded_box(draw, boxes["webhooks"], "payos_webhook_events", ["orderCode / reference", "signatureHash", "processingStatus", "dormant"], accent="#8B5E83")
    rounded_box(draw, boxes["counters"], "counters", ["_id: payos_order_code", "sequenceValue", "atomic increment"], accent="#667085")

    arrow(draw, (440, 280), (590, 280), "1:N")
    arrow(draw, (260, 410), (260, 520), "1:N")
    arrow(draw, (760, 420), (720, 540), "1:N")
    arrow(draw, (880, 420), (1080, 540), "1:N")
    poly_arrow(draw, [(440, 220), (440, 155), (1710, 155), (1710, 180)], "1:N", (1080, 145))
    arrow(draw, (1710, 455), (1570, 560), "1:N")
    arrow(draw, (1850, 455), (1930, 560), "1:N")
    arrow(draw, (1570, 825), (1570, 960), "1:N")
    arrow(draw, (1750, 700), (1790, 700), "1:N")
    poly_arrow(draw, [(1930, 865), (1930, 915), (1750, 1040)], "0..1:1", (1870, 910))
    poly_arrow(draw, [(1450, 350), (1320, 350), (1320, 930), (1110, 980)], "1:N", (1305, 720))
    draw.text((150, 1300), "Liên kết mềm: counters cấp orderCode; payos_webhook_events đối chiếu hai payment collection bằng orderCode.", font=pil_font(22, True), fill="#667085")

    draw.rounded_rectangle((70, 1360, 2030, 1450), radius=20, fill="#FFF7E6", outline="#E3B65A", width=3)
    draw.text((95, 1382), "Lưu ý: saving_events và payos_webhook_events đã có schema nhưng service chưa ghi; init script hiện thiếu push_subscriptions.", font=pil_font(25, True), fill="#7A5A00")
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=95)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_widths(column_count: int) -> list[int]:
    presets = {
        2: [2700, 6660],
        3: [1800, 2700, 4860],
        4: [1200, 3000, 1800, 3360],
        5: [760, 2820, 760, 2360, 2660],
    }
    return presets.get(column_count, [TABLE_WIDTH_DXA // column_count] * column_count)


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, size: float | None = None, color: str | None = None, bold: bool = False) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(size or 11) - 0.5, color=DARK_BLUE, bold=False)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color, bold=bold)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 30, INK, 0, 8),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=MUTED)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("SAVING 100  |  TÀI LIỆU KỸ THUẬT")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(p)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("HỒ SƠ KIẾN TRÚC VÀ ĐẶC TẢ HỆ THỐNG")
    set_run_font(run, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("SAVING 100")
    set_run_font(run, size=34, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("Database · UI · Backend · API · Auth · payOS · Web Push")
    set_run_font(run, size=14, color=MUTED)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(30)
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, size=10, color=BLUE)

    metadata = [
        ("Phiên bản", "1.0"),
        ("Ngày rà soát", "30/07/2026"),
        ("Mốc source", "764698d037af2297f3cc5177bcf4ee5557de987c"),
        ("Trạng thái", "Tài liệu hiện trạng (as-is)"),
        ("Package manager", "pnpm@9.15.0"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(metadata):
        left, right = table.rows[index].cells
        left.text = ""
        right.text = ""
        p_left = left.paragraphs[0]
        add_inline(p_left, label, size=10, color=MUTED, bold=True)
        p_right = right.paragraphs[0]
        add_inline(p_right, value, size=10, color=BLACK)
        set_cell_shading(left, LIGHT_GRAY)
    set_table_geometry(table, [2300, 7060])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tài liệu không chứa secret, mật khẩu hoặc khóa tích hợp thực tế.")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    p = doc.add_paragraph("MỤC LỤC", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    sections = [
        "1. Tổng quan hệ thống",
        "2. Phạm vi chức năng hiện có",
        "3. Kiến trúc tổng thể",
        "4. Cấu trúc source và công nghệ",
        "5. Thiết kế frontend và UI",
        "6. Thiết kế backend",
        "7. Mô hình dữ liệu và ERD",
        "8. Quy tắc nghiệp vụ và vòng đời trạng thái",
        "9. Đặc tả API REST",
        "10. Xác thực, cookie và bảo mật",
        "11. Tích hợp payOS",
        "12. Web Push và lịch nhắc",
        "13. Cấu hình môi trường và triển khai",
        "14. Kiểm thử, chất lượng và vận hành",
        "15. Sai lệch, rủi ro và lộ trình chuẩn hoá",
        "16. Thuật ngữ và nguồn đối chiếu",
    ]
    for item in sections:
        p = doc.add_paragraph(style="List Number")
        # The source already contains section numbers; use a regular paragraph to avoid double numbering.
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        add_inline(p, item, size=10.5, color=DARK_BLUE, bold=True)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    add_inline(note, "Phạm vi: source hiện tại tại ngày 30/07/2026. Các điểm chưa thực thi được đánh dấu rõ trong từng mục.", size=10, color=MUTED)
    doc.add_page_break()


def add_image(doc: Document, alt: str, relative_path: str) -> None:
    path = ROOT / relative_path
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", alt)
        doc_pr[0].set("title", alt)
    caption = doc.add_paragraph(alt, style="Caption")
    caption.paragraph_format.keep_with_next = False


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    for line_index, line in enumerate(code.splitlines()):
        if line_index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=BLACK)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text, size=10, color=DARK_BLUE)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            rows.append(parts)
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=columns)
    table.style = "Table Grid"
    for row_index, row_data in enumerate(normalized):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline(p, value, size=8.5 if columns >= 4 else 9, color=BLACK, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
        if row_index == 0:
            set_repeat_table_header(table.rows[row_index])
    set_table_geometry(table, table_widths(columns))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("# 1. "))
    lines = lines[start:]
    index = 0
    first_h1 = True
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("```"):
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(doc, "\n".join(code_lines))
            continue
        if line.startswith("!["):
            match = re.fullmatch(r"!\[([^\]]+)\]\(([^)]+)\)", line)
            if match:
                add_image(doc, match.group(1), match.group(2))
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if line.startswith(">"):
            add_callout(doc, line.lstrip("> ").strip())
            index += 1
            continue
        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            p.paragraph_format.page_break_before = False
            index += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith(("#", "|", "```", "![", ">", "- "))
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    create_architecture_image(ASSETS / "architecture-overview.png")
    create_erd_image(ASSETS / "erd-current.png")

    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_static_toc(doc)
    add_markdown_body(doc, SOURCE.read_text(encoding="utf-8"))
    set_update_fields(doc)

    props = doc.core_properties
    props.title = "Tài liệu kỹ thuật hệ thống Saving 100"
    props.subject = "Database, UI, backend, API, auth, payOS và Web Push"
    props.author = "Saving 100 / Codex"
    props.keywords = "Saving 100, NestJS, Next.js, MongoDB, payOS, Web Push, ERD, API"
    props.comments = "Generated from the reviewed source at commit 764698d037af2297f3cc5177bcf4ee5557de987c."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
